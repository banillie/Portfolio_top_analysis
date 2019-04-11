'''Programme produces outputs (figures) for building the speeddials page

outputs:
programme prints out high level numbers for speeddials
overall SRO DCA changes between quarters - in MS word. Typically current and last. But can be adapted to whichever two
quarters are of interest
overall SRO finance confidence changes between quarters - as above
overall SRO benefits confidence changes between quarters - as above
overall SRO resource confidence changes between quarters - as above

Operating requirements:
1) Specify file path to core data sets of interest.
2) specify file path and file name for output MS word documents. '''


from bcompiler.utils import project_data_from_master
import docx


def get_project_names(dictionary):
    project_name_list = []
    for x in dictionary:
        project_name_list.append(x)
    return project_name_list


def create_dict(category, dictionary_1, dictionary_2):
    m_dict = {}
    for name in dictionary_1:
        m_dict[name] = {}
        holding_dict_1 = dictionary_1[name]
        a = holding_dict_1[category]
        m_dict[name]['DCA'] = a
        try:
            holding_dict_2 = dictionary_2[name]
            b = holding_dict_2[category]
            m_dict[name]['DCA_lq'] = b
        except KeyError:
            b = 'not reporting'
            m_dict[name]['DCA_lq'] = b
        if a == b:
            m_dict[name]['Change'] = 0
        elif a != b:
            if m_dict[name]['DCA_lq'] == 'not reporting':
                m_dict[name]['Change'] = 0
            elif m_dict[name]['DCA'] == None:  # quick hack fix here for handling projects that remove dca reporting
                m_dict[name]['Change'] = 0
            elif m_dict[name]['DCA'] == 'Green':
                if m_dict[name]['DCA_lq'] == 'Amber/Green':
                    m_dict[name]['Change'] = 1
                if m_dict[name]['DCA_lq'] == 'Amber':
                    m_dict[name]['Change'] = 2
                if m_dict[name]['DCA_lq'] == 'Amber/Red':
                    m_dict[name]['Change'] = 3
                if m_dict[name]['DCA_lq'] == 'Red':
                    m_dict[name]['Change'] = 4
                if m_dict[name]['DCA_lq'] == None:
                    m_dict[name]['Change'] = 5
            elif m_dict[name]['DCA'] == 'Amber/Green':
                if m_dict[name]['DCA_lq'] == 'Green':
                    m_dict[name]['Change'] = -1
                if m_dict[name]['DCA_lq'] == 'Amber':
                    m_dict[name]['Change'] = 1
                if m_dict[name]['DCA_lq'] == 'Amber/Red':
                    m_dict[name]['Change'] = 2
                if m_dict[name]['DCA_lq'] == 'Red':
                    m_dict[name]['Change'] = 3
                if m_dict[name]['DCA_lq'] == None:
                    m_dict[name]['Change'] = 5
            elif m_dict[name]['DCA'] == 'Amber':
                if m_dict[name]['DCA_lq'] == 'Green':
                    m_dict[name]['Change'] = -2
                if m_dict[name]['DCA_lq'] == 'Amber/Green':
                    m_dict[name]['Change'] = -1
                if m_dict[name]['DCA_lq'] == 'Amber/Red':
                    m_dict[name]['Change'] = 1
                if m_dict[name]['DCA_lq'] == 'Red':
                    m_dict[name]['Change'] = 2
                if m_dict[name]['DCA_lq'] == None:
                    m_dict[name]['Change'] = 5
            elif m_dict[name]['DCA'] == 'Amber/Red':
                if m_dict[name]['DCA_lq'] == 'Green':
                    m_dict[name]['Change'] = -3
                if m_dict[name]['DCA_lq'] == 'Amber/Green':
                    m_dict[name]['Change'] = -2
                if m_dict[name]['DCA_lq'] == 'Amber':
                    m_dict[name]['Change'] = -1
                if m_dict[name]['DCA_lq'] == 'Red':
                    m_dict[name]['Change'] = 1
                if m_dict[name]['DCA_lq'] == None:
                    m_dict[name]['Change'] = 5
            elif m_dict[name]['DCA'] == 'Red':
                if m_dict[name]['DCA_lq'] == 'Green':
                    m_dict[name]['Change'] = -4
                if m_dict[name]['DCA_lq'] == 'Amber/Green':
                    m_dict[name]['Change'] = -3
                if m_dict[name]['DCA_lq'] == 'Amber':
                    m_dict[name]['Change'] = -2
                if m_dict[name]['DCA_lq'] == 'Amber/Red':
                    m_dict[name]['Change'] = -1
                if m_dict[name]['DCA_lq'] == None:
                    m_dict[name]['Change'] = 5

    return m_dict


def printing_word(m_dict):
    doc = docx.Document()
    title = 'Confidence changes this quarter'
    top = doc.add_paragraph()
    top.add_run(title).bold = True
    doc.add_paragraph()
    sub_head = 'Decrease (in order of size of change)'
    sub = doc.add_paragraph()
    sub.add_run(sub_head).bold = True
    down = 0
    for name in m_dict:
        # print(name, m_dict[name])
        if m_dict[name]['Change'] == -4:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            c = str(name)
            down += 1
            d = str(down) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    for name in m_dict:
        if m_dict[name]['Change'] == -3:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            c = str(name)
            down += 1
            d = str(down) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    for name in m_dict:
        if m_dict[name]['Change'] == -2:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            c = str(name)
            down += 1
            d = str(down) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    for name in m_dict:
        if m_dict[name]['Change'] == -1:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            c = str(name)
            down += 1
            d = str(down) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    doc.add_paragraph()
    g = doc.add_paragraph()
    totals_line = str(down) + ' project(s) have decreased in total'
    g.add_run(totals_line).bold = True
    doc.add_paragraph()
    sub_head_2 = 'Increase (in order of size of change)'
    sub_2 = doc.add_paragraph()
    sub_2.add_run(sub_head_2).bold = True
    up = 0
    for name in m_dict:
        if m_dict[name]['Change'] == 4:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            c = str(name)
            up += 1
            d = str(up) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    for name in m_dict:
        if m_dict[name]['Change'] == 3:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            c = str(name)
            up += 1
            d = str(up) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    for name in m_dict:
        if m_dict[name]['Change'] == 2:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            up += 1
            c = str(name)
            d = str(up) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    for name in m_dict:
        if m_dict[name]['Change'] == 1:
            p = doc.add_paragraph()
            a = m_dict[name]['DCA']
            b = m_dict[name]['DCA_lq']
            c = str(name)
            up += 1
            d = str(up) + '. ' + c
            e = ': change from ' + str(b) + ' to ' + str(a)
            p.add_run(d).bold = True
            p.add_run(e)
    doc.add_paragraph()
    totals_line_2 = str(up) + ' project(s) have increased in total'
    h = doc.add_paragraph()
    h.add_run(totals_line_2).bold = True

    return doc


def calculating_dials_dca(dictionary):
    empty_list = []
    for x in dictionary:
        empty_list.append(dictionary[x]['DCA'])

    count_list = []
    Red = empty_list.count('Red')
    count_list.append(('Red', Red))
    amber_red = empty_list.count('Amber/Red')
    count_list.append(('Amber/Red', amber_red))
    amber = empty_list.count('Amber')
    count_list.append(('Amber', amber))
    amber_green = empty_list.count('Amber/Green')
    count_list.append(('Amber/Green', amber_green))
    green = empty_list.count('Green')
    count_list.append(('Green', green))

    print(count_list)

    total = 0
    for i in range(0, len(count_list)):
        total += (count_list[i][1])

    print('total number of projects ' + str(total))

    a = count_list[0][1] * 0
    b = count_list[1][1] * 25
    c = count_list[2][1] * 50
    d = count_list[3][1] * 75
    e = count_list[4][1] * 100

    score = a + b + c + d + e
    maximum = total * 100

    result = score / maximum

    print(result)


def calculating_dials_other(dictionary):
    empty_list = []
    for x in dictionary:
        empty_list.append(dictionary[x]['DCA'])

    count_list = []
    Red = empty_list.count('Red')
    count_list.append(('Red', Red))
    amber = empty_list.count('Amber')
    count_list.append(('Amber', amber))
    green = empty_list.count('Green')
    count_list.append(('Green', green))

    print(count_list)

    total = 0
    for i in range(0, len(count_list)):
        total += (count_list[i][1])

    print('total number of projects ' + str(total))

    a = count_list[0][1] * 0
    b = count_list[1][1] * 50
    c = count_list[2][1] * 100

    score = a + b + c
    maximum = total * 100

    result = score / maximum

    print(result)


# 1) Specify file path to core data sets of interest.

current_Q_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\masters folder\\core data\\merged_master_'
                                          'testing.xlsx')

last_Q_dict = project_data_from_master('C:\\Users\\Standalone\\Will\\masters folder\\core data\\master_2_2018.xlsx')

current_Q_list = get_project_names(current_Q_dict)

sro_dca = create_dict('Departmental DCA', current_Q_dict, last_Q_dict)
finance_dca = create_dict('SRO Finance confidence', current_Q_dict, last_Q_dict)
resource_dca = create_dict('Overall Resource DCA - Now', current_Q_dict, last_Q_dict)
benefits_dca = create_dict('SRO Benefits RAG', current_Q_dict, last_Q_dict)

print('DCA')
calculating_dials_dca(sro_dca)
print('Finance')
calculating_dials_other(finance_dca)
print('Resource')
calculating_dials_other(resource_dca)
print('Benefits')
calculating_dials_other(benefits_dca)

overall = printing_word(sro_dca)
finance = printing_word(finance_dca)
resource = printing_word(resource_dca)
benefits = printing_word(benefits_dca)

# 2) specify file path and file name for output MS word documents.
overall.save('C:\\Users\\Standalone\\Will\\Q3_1718_overall_dca.docx')
finance.save('C:\\Users\\Standalone\\Will\\Q3_1819_finance_dca.docx')
resource.save('C:\\Users\\Standalone\\Will\\Q3_1819_resource_dca.docx')
benefits.save('C:\\Users\\Standalone\\Will\\Q3_1819_benefits_dca.docx')
